# -*- coding: utf-8 -*-
"""
filldoc/generator.py

Совместим с вызовом из твоего cli:
    apply_mapping_to_doc(d, data_row, mapping_rows, settings, log=log, dry_run=dry_run)

Поддерживает selector_type:
- number           — замена цифро-меток: 1, (1), [1], «1», "1"
- near_label       — поиск подписи в таблицах и запись right/below
- cell             — адресная подстановка: table=0,row=3,col=1
(опционально)
- regex            — замена по регулярному выражению внутри параграфов/ячеек
- table_header_cell— поиск столбца по заголовку первой строки таблицы

Поддерживает fmt:
- date(DD.MM.YYYY), upper, lower, initials, trim, default("—")
"""

from __future__ import annotations
import re
from typing import Iterable, List, Dict, Any, Optional
from datetime import datetime, date

from docx.document import Document as DocxDocument
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


# ---------------------- УТИЛИТЫ ФОРМАТОВ ---------------------- #

def _py_datefmt(mask: str) -> str:
    # Excel-подобные маски -> strftime
    return (mask.replace("DD", "%d")
                .replace("MM", "%m")
                .replace("YYYY", "%Y"))

def _fmt_value(value: Any, fmt: Optional[str], settings: Dict[str, Any]) -> str:
    if value is None:
        value = ""

    s = str(value)

    if not fmt:
        return s

    fmt_low = fmt.strip().lower()

    # date(DD.MM.YYYY)
    m = re.fullmatch(r"date\((.+?)\)", fmt.strip(), flags=re.I)
    if m:
        out_mask = _py_datefmt(m.group(1))
        # Попытка умного парсинга входного значения
        if isinstance(value, (date, datetime)):
            return value.strftime(out_mask)
        for pat in ("%Y-%m-%d", "%Y/%m/%d", "%d.%m.%Y", "%d/%m/%Y"):
            try:
                dt = datetime.strptime(s, pat)
                return dt.strftime(out_mask)
            except Exception:
                pass
        # Безуспешно — оставляем как есть
        return s

    # default("—")
    m = re.fullmatch(r'default\("(.*)"\)', fmt.strip())
    if m:
        return s if s.strip() else m.group(1)

    if fmt_low == "upper":
        return s.upper()
    if fmt_low == "lower":
        return s.lower()
    if fmt_low == "trim":
        return s.strip()
    if fmt_low == "initials":
        # "Иванов Иван Иванович" -> "Иванов И.И."
        parts = s.split()
        if not parts:
            return s
        last = parts[0]
        initials = "".join((p[0] + ".") for p in parts[1:] if p)
        return f"{last} {initials}"

    # Неизвестный fmt — вернём как есть
    return s


# ---------------------- ОБХОД ДОКУМЕНТА ---------------------- #

def _all_paragraphs(doc: DocxDocument) -> Iterable[Paragraph]:
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p

def _set_paragraph_text(par: Paragraph, text: str) -> None:
    # Полная замена текста параграфа с очисткой runs
    # (теряется локальная разметка — это ожидаемо для автоподстановок)
    # Сохраняем только первый run и пишем в него
    for run in list(par.runs)[1:]:
        r_el = run._element
        r_el.getparent().remove(r_el)
    if par.runs:
        par.runs[0].text = text
    else:
        par.add_run(text)

def _set_cell_text(cell: _Cell, text: str) -> None:
    # Очистить все параграфы в ячейке и записать одну строку
    for p in list(cell.paragraphs)[1:]:
        p._element.getparent().remove(p._element)
    if cell.paragraphs:
        _set_paragraph_text(cell.paragraphs[0], text)
    else:
        p = cell.add_paragraph("")
        _set_paragraph_text(p, text)

def _find_label_cell(table: Table, label: str) -> Optional[tuple[int, int]]:
    # Ищем точное совпадение текста ячейки (с Trim)
    target = label.strip()
    target_no_colon = target[:-1].strip() if target.endswith(":") else None
    for r_i, row in enumerate(table.rows):
        for c_i, cell in enumerate(row.cells):
            txt = "\n".join(p.text for p in cell.paragraphs).strip()
            if txt == target:
                return (r_i, c_i)
            # Иногда подпись в Excel указана без двоеточия, а в шаблоне с двоеточием (или наоборот)
            if target_no_colon and txt == target_no_colon:
                return (r_i, c_i)
            if target.endswith(":") and txt.rstrip() == target[:-1]:
                return (r_i, c_i)
    return None


# ---------------------- SELECTOR HELPERS ---------------------- #

def _number_patterns(num: str) -> List[re.Pattern]:
    # Поддерживаем: 1, (1), [1], «1», "1"
    escaped = re.escape(num)
    return [
        re.compile(rf"\b{escaped}\b"),
        re.compile(rf"\({escaped}\)"),
        re.compile(rf"\[{escaped}\]"),
        re.compile(rf"[«\"]{escaped}[»\"]"),
    ]

def _replace_in_runs(par: Paragraph, pattern: re.Pattern, repl: str) -> int:
    count = 0
    for run in par.runs:
        new_text, n = pattern.subn(repl, run.text)
        if n:
            run.text = new_text
            count += n
    return count

def _parse_cell_selector(selector: str) -> tuple[int, int, int]:
    """
    selector: "table=0,row=3,col=1"
    Возвращает кортеж индексов (t_i, r_i, c_i), нумерация с 0.
    """
    parts = dict(
        (k.strip(), v.strip())
        for k, v in (p.split("=", 1) for p in selector.split(","))
    )
    return int(parts["table"]), int(parts["row"]), int(parts["col"])


# ---------------------- УСЛОВИЯ when ---------------------- #

def _match_when(expr: str, data_row: Dict[str, Any]) -> bool:
    """
    Примитивный парсер условий:
      - "Курс=3"
      - "Группа like ИВТ-*"
      - несколько условий через ' and ' (пробел-and-пробел)
    Любые ошибки парсинга → True (не блокируем правило).
    """
    if not expr or not str(expr).strip():
        return True

    def _like(pattern: str, value: str) -> bool:
        # поддерживаем * как любой хвост
        # простая проверка: prefix*  /  *suffix  /  *mid*
        p = pattern
        v = value or ""
        if p.startswith("*") and p.endswith("*"):
            return p.strip("*") in v
        if p.startswith("*"):
            return v.endswith(p[1:])
        if p.endswith("*"):
            return v.startswith(p[:-1])
        return v == p

    try:
        chunks = [c.strip() for c in str(expr).split(" and ")]
        for chunk in chunks:
            if " like " in chunk:
                key, pat = [x.strip() for x in chunk.split(" like ", 1)]
                if _like(pat, str(data_row.get(key, ""))) is False:
                    return False
            elif "!=" in chunk:
                key, val = [x.strip() for x in chunk.split("!=", 1)]
                if str(data_row.get(key, "")) == val:
                    return False
            elif "=" in chunk:
                key, val = [x.strip() for x in chunk.split("=", 1)]
                if str(data_row.get(key, "")) != val:
                    return False
            else:
                # неизвестный синтаксис — не блокируем
                pass
        return True
    except Exception:
        return True


# ---------------------- ОСНОВНОЙ ДВИЖОК ---------------------- #

def apply_mapping_to_doc(doc: DocxDocument,
                         data_row: Dict[str, Any],
                         mapping_rows: List[Dict[str, Any]],
                         settings: Dict[str, Any],
                         log: Optional[List[str]] = None,
                         dry_run: bool = False) -> None:
    """
    mapping_rows: список правил (dict) со столбцами:
      - doc_type (уже отфильтрован на этапе CLI/спека)
      - selector_type: number | near_label | cell | regex | table_header_cell
      - selector: строка селектора
      - direction: right | below  (для near_label)
      - key: имя поля из data_row
      - fmt: формат отображения
      - when: условие применения (опц.)
    """
    if log is None:
        log = []

    # Глобальные настройки (на будущее)
    # date_fmt_default = settings.get("date_format")  # пока не используется напрямую

    for rule in mapping_rows:
        stype = str(rule.get("selector_type", "")).strip().lower()
        selector = str(rule.get("selector", "")).strip()
        key = str(rule.get("key", "")).strip()
        fmt = str(rule.get("fmt", "")).strip()
        direction = str(rule.get("direction", "")).strip().lower()
        cond = str(rule.get("when", "")).strip()

        # Базовая валидация строки
        if not stype or not selector:
            log.append("SKIP: пустой selector_type/selector")
            continue

        # Проверка when
        if not _match_when(cond, data_row):
            log.append(f"SKIP by when: {cond}")
            continue

        # Получаем значение
        value = _fmt_value(data_row.get(key), fmt, settings)

        # ---- number ----
        if stype == "number":
            if not selector:
                log.append("number: пустой selector (пропуск)")
                continue
            total = 0
            for pat in _number_patterns(selector):
                for p in _all_paragraphs(doc):
                    total += _replace_in_runs(p, pat, value)
            log.append(f"number {selector} -> '{value}' ({total})")

        # ---- near_label ----
        elif stype == "near_label":
            if direction not in ("right", "below"):
                log.append(f"near_label '{selector}' -> SKIP (нужен direction=right|below)")
                continue
            placed = False
            for t_i, tbl in enumerate(doc.tables):
                pos = _find_label_cell(tbl, selector)
                if not pos:
                    continue
                r_i, c_i = pos
                try:
                    if direction == "right":
                        target = tbl.cell(r_i, c_i + 1)
                    else:  # below
                        target = tbl.cell(r_i + 1, c_i)
                except Exception:
                    log.append(f"near_label '{selector}' {direction} -> SKIP (нет целевой ячейки)")
                    continue

                if not dry_run:
                    _set_cell_text(target, value)
                log.append(f"near_label '{selector}' {direction} -> '{value}' (table {t_i})")
                placed = True
                break
            if not placed:
                log.append(f"near_label '{selector}' -> не найдено")

        # ---- cell ----
        elif stype == "cell":
            try:
                t_i, r_i, c_i = _parse_cell_selector(selector)
                cell = doc.tables[t_i].rows[r_i].cells[c_i]
                if not dry_run:
                    _set_cell_text(cell, value)
                log.append(f"cell {selector} -> '{value}'")
            except Exception:
                log.append(f"cell {selector} -> SKIP (нет такой ячейки)")

        # ---- regex (опционально) ----
        elif stype == "regex":
            try:
                pat = re.compile(selector)
            except re.error as e:
                log.append(f"regex /{selector}/ -> SKIP (ошибка компиляции: {e})")
                continue
            total = 0
            for p in _all_paragraphs(doc):
                # regex может пересекать runs — поэтому заменяем весь параграф целиком
                new_text, n = pat.subn(value, p.text)
                if n:
                    total += n
                    if not dry_run:
                        _set_paragraph_text(p, new_text)
            log.append(f"regex /{selector}/ -> '{value}' ({total})")

        # ---- table_header_cell (опционально) ----
        elif stype == "table_header_cell":
            # selector: "table=0,header='ФИО',row=1"
            m = re.findall(r"(\w+)=('([^']+)'|\"([^\"]+)\"|[^,]+)", selector)
            params = {k: (v2 or v3 or v).strip("'\"") for k, v, v2, v3 in m}
            try:
                t_i = int(params.get("table", "0"))
                header = params["header"]
                row_index = int(params.get("row", "1"))
                tbl = doc.tables[t_i]
                # найдем столбец, где первая строка == header
                col_idx = None
                if tbl.rows:
                    for c_i, cell in enumerate(tbl.rows[0].cells):
                        if "\n".join(p.text for p in cell.paragraphs).strip() == header:
                            col_idx = c_i
                            break
                if col_idx is None:
                    log.append(f"table_header_cell {selector} -> SKIP (нет столбца '{header}')")
                else:
                    cell = tbl.cell(row_index, col_idx)
                    if not dry_run:
                        _set_cell_text(cell, value)
                    log.append(f"table_header_cell {selector} -> '{value}'")
            except Exception:
                log.append(f"table_header_cell {selector} -> SKIP (ошибка параметров)")

        else:
            log.append(f"Unknown selector_type: {stype}")

    # Функция ничего не возвращает; все изменения внесены в doc (если not dry_run),
    # а подробный ход записан в log.
