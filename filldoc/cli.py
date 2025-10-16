# filldoc/cli.py
# -*- coding: utf-8 -*-
import argparse, os, glob, re
from docx import Document
from .utils import convert_doc_to_docx
from .detector import detect_template, suggest_anchors
from .templates import TEMPLATES
from fnmatch import fnmatchcase
from .generator import apply_mapping_to_doc
from pathlib import Path
from fnmatch import fnmatchcase
from docx import Document
import pandas as pd
import inspect
import inspect as _insp

try:
    from .cli_single import run_single
except ImportError:
    def run_single(args):
        raise SystemExit("Отсутствует filldoc/cli_single.py — режим --doc/--excel недоступен.")

def render_mask(mask: str, row: dict) -> str:
    def _repl(m):
        key = m.group(1).strip()
        val = row.get(key)
        return "" if val is None else str(val)
    return re.sub(r"\{\{\s*([^}]+?)\s*\}\}", _repl, mask)

def _lazy_import_run_single():
    try:
        from .cli_single import run_single
        return run_single
    except Exception as e:
        import sys
        print("Для режима --auto нужен файл filldoc/cli_single.py с функцией run_single(...). "
              "Сейчас он не найден.\nДетали:", e, file=sys.stderr)
        sys.exit(1)

def _process_detect(path: str):
    try:
        src = convert_doc_to_docx(path)
    except Exception as e:
        print(f"[CONVERT-ERROR] {os.path.basename(path)}: {e}")
        return
    doc = Document(src)
    det = detect_template(doc)
    base = os.path.basename(path)
    if det.template_id:
        spec = TEMPLATES[det.template_id]
        print(f"[DETECT] '{base}' → {det.template_id} ({spec.human_name}), score={det.score}")
        for note in det.breakdown:
            print("  -", note)
    else:
        print(f"[DETECT] '{base}' → НЕ ОПОЗНАН (score={det.score})")
        for note in det.breakdown:
            print("  -", note)

def _process_suggest(path: str):
    try:
        src = convert_doc_to_docx(path)
    except Exception as e:
        print(f"[CONVERT-ERROR] {os.path.basename(path)}: {e}")
        return
    doc = Document(src)
    base = os.path.basename(path)
    hints = suggest_anchors(doc)
    print(f"\n[SUGGEST] {base}")
    print("  top_lines:")
    for s in hints["top_lines"]:
        print("   •", s)
    print("  with_colon:")
    for s in hints["with_colon"]:
        print("   •", s)
    print("  uppercase:")
    for s in hints["uppercase"]:
        print("   •", s)
    print("  keywords:")
    for s in hints["keywords"]:
        print("   •", s)
    # рыба для вставки в templates.py
    print("  --- TemplateSpec skeleton ---")
    print("  required = [")
    for s in (hints["uppercase"][:2] + hints["with_colon"][:3] + hints["keywords"][:1]):
        print(f"    \"{s}\",")
    print("  ]")
    print("  optional = [")
    for s in (hints["with_colon"][3:6] + hints["top_lines"][2:4]):
        print(f"    \"{s}\",")
    print("  ]")

def _process_auto(src: str, excel: str, row: int | None, out_dir: str, dry_run: bool) -> None:
    """
    Авто-режим:
    1) Загружает данные из Excel (лист 'data').
    2) Для каждого входного DOCX (файл или все .docx из папки):
       - сперва пытается подобрать шаблон по имени файла (settings['filename_globs']);
       - если несколько совпало — сужает старым контент-детектором;
       - если ничего — использует старый контент-детектор;
    3) Применяет mapping выбранного шаблона к документу с данными заданной строки Excel
       (или ко всем строкам, если row не указан);
    4) Печатает предпросмотр (dry-run) или сохраняет готовый файл.
    """
    import os
    from pathlib import Path
    from fnmatch import fnmatchcase
    from docx import Document
    import pandas as pd
    import inspect

    # ---------- helpers ----------
    def _pick_by_filename(path: str, templates):
        name = os.path.basename(path).lower()
        winners = []
        for t in templates:
            globs = (t.settings or {}).get("filename_globs", []) or (t.settings or {}).get("file_globs", [])
            for g in globs:
                if fnmatchcase(name, g.lower()):
                    winners.append(t)
                    break
        return winners

    def _render_mask(mask: str, data: dict, fallback_name: str) -> str:
        """Наивная подстановка {{Ключ}} значениями data (None -> '')."""
        if not mask:
            return fallback_name
        out = mask
        for k, v in data.items():
            out = out.replace(f"{{{{{k}}}}}", "" if v is None else str(v))
        return out

    # ---------- входные документы ----------
    if not src:
        raise SystemExit("--auto: не указан путь (--doc/--input).")

    p = Path(src)
    if p.is_dir():
        paths = sorted(q for q in p.iterdir() if q.is_file() and q.suffix.lower() in (".docx", ".doc"))
    else:
        if not p.exists():
            raise SystemExit(f"Файл не найден: {src}")
        if p.suffix.lower() not in (".docx", ".doc"):
            print(f"[AUTO] Пропуск '{p.name}': неподдерживаемое расширение")
            return
        paths = [p]

    # ---------- Excel: лист 'data' ----------
    try:
        df = pd.read_excel(excel, sheet_name="data")
    except Exception as e:
        raise SystemExit(f"Не удалось прочитать лист 'data' из {excel}: {e}")
    if df.empty:
        raise SystemExit(f"В {excel} на листе 'data' нет строк.")

    # строки данных (row — 1-базный)
    if row is not None:
        idx = row - 1
        if idx < 0 or idx >= len(df):
            raise SystemExit(f"--row {row}: вне диапазона (в 'data' всего {len(df)} строк).")
        rows = [df.iloc[idx].to_dict()]
    else:
        rows = df.to_dict(orient="records")

    # ---------- целевая папка ----------
    Path(out_dir).mkdir(parents=True, exist_ok=True)

    # ---------- основной цикл ----------
    for orig_path in paths:
        base = orig_path.name
        # открываем для детекта
        try:
            doc_for_detect = Document(str(orig_path))
        except Exception as e:
            print(f"[AUTO] Пропуск '{base}': не удалось открыть DOCX ({e})")
            continue

        templates_all = list(TEMPLATES.values())

        # 1) приоритет — по имени файла
        spec = None
        det = None
        by_name = _pick_by_filename(str(orig_path), templates_all)  # ИСПОЛЬЗУЕМ ИМЯ ИСХОДНИКА
        if len(by_name) == 1:
            spec = by_name[0]
            det = type("Det", (), {"template_id": spec.id, "score": 999})()
        elif len(by_name) > 1:
            try:
                det = detect_template(doc_for_detect, candidates=by_name)  # type: ignore[arg-type]
                spec = TEMPLATES.get(det.template_id) if det and getattr(det, "template_id", None) else by_name[0]
            except TypeError:
                det = detect_template(doc_for_detect)
                spec = TEMPLATES.get(det.template_id) if det and getattr(det, "template_id", None) else by_name[0]
        else:
            # 2) фоллбэк — контент-детект
            det = detect_template(doc_for_detect)
            spec = TEMPLATES.get(det.template_id) if det and getattr(det, "template_id", None) else None

        if not spec:
            print(f"[AUTO] '{base}' → не опознан (score={getattr(det,'score',0)})")
            continue

        # ---------- применяем mapping для выбранных строк Excel ----------
        for r in rows:
            # подготовим данные + {{SRC}}
            r_with_src = dict(r)
            r_with_src["SRC"] = orig_path.stem

            # имя файла
            name_mask = (spec.settings or {}).get("маска_имени_файла") or base
            out_name = _render_mask(name_mask, r_with_src, base)
            out_file = Path(out_dir, out_name)

            # разводим коллизии имён
            if out_file.exists():
                stem, suf = out_file.stem, out_file.suffix
                k = 2
                while True:
                    cand = out_file.with_name(f"{stem}_{k}{suf}")
                    if not cand.exists():
                        out_file = cand
                        break
                    k += 1
            out_path = str(out_file)

            # открыть «чистую» копию для подстановки
            try:
                d = Document(str(orig_path))
            except Exception as e:
                print(f"[{spec.id}] Пропуск '{base}': не удалось открыть повторно ({e})")
                continue

            # безопасно вызвать apply_mapping_to_doc (проверка на наличие log)
            log = []
            called = False
            err = None
            try:
                sig = inspect.signature(apply_mapping_to_doc)
                params = sig.parameters
                if "log" in params and "dry_run" in params:
                    print(f"[{spec.id}] Применяем mapping: {spec.mapping}")
                    apply_mapping_to_doc(d, r_with_src, spec.mapping, spec.settings, log=log, dry_run=dry_run)
                    called = True
                elif "dry_run" in params and "log" not in params:
                    print(f"[{spec.id}] Применяем mapping: {spec.mapping}")
                    apply_mapping_to_doc(d, r_with_src, spec.mapping, spec.settings, dry_run=dry_run)
                    called = True
                elif "log" in params and "dry_run" not in params:
                    print(f"[{spec.id}] Применяем mapping: {spec.mapping}")
                    apply_mapping_to_doc(d, r_with_src, spec.mapping, spec.settings, log)
                    called = True
                else:
                    print(f"[{spec.id}] Применяем mapping: {spec.mapping}")
                    apply_mapping_to_doc(d, r_with_src, spec.mapping, spec.settings)
                    called = True
            except TypeError as e:
                err = e
            except Exception as e:
                err = e

            if not called:
                print(f"[{spec.id}] Ошибка применения mapping к '{base}': {err}")
                continue

            # предпросмотр / сохранение
            if dry_run:
                print(f"[{spec.id}] Предпросмотр: {out_file.name}")
                if log:
                    for line in log:
                        print("  -", line)
                else:
                    print("  (без подробного лога: сигнатура apply_mapping_to_doc не принимает 'log')")
            else:
                try:
                    d.save(out_path)
                    print(f"[{spec.id}] Сохранено: {out_path}")
                except Exception as e:
                    print(f"[{spec.id}] Не удалось сохранить '{out_path}': {e}")


def main():
    p = argparse.ArgumentParser()
    p.add_argument("--detect", action="store_true", help="распознать шаблон по 1-й странице")
    p.add_argument("--suggest", action="store_true", help="подсказать якоря по 1-й странице (для заполнения templates.py)")
    p.add_argument("--auto", action="store_true", help="распознать и заполнить")
    p.add_argument("--doc", help="входной .docx|.doc")
    p.add_argument("--input", help="папка с документами")
    p.add_argument("--excel", help="master.xlsx с листом data (для --auto)")
    p.add_argument("--row", type=int)
    p.add_argument("--out", default="./output")
    p.add_argument("--dry-run", action="store_true")
    args = p.parse_args()

    if args.detect or args.suggest:
        handler = _process_suggest if args.suggest else _process_detect
        if args.input:
            for f in glob.glob(os.path.join(args.input, "*.*")):
                if f.lower().endswith((".docx", ".doc")):
                    handler(f)
        else:
            if not args.doc:
                raise SystemExit("Укажите --doc или --input")
            handler(args.doc)
        return

    if args.auto:
        if not args.excel:
            raise SystemExit("Для --auto нужен --excel (лист 'data').")
        if args.input:
            for f in glob.glob(os.path.join(args.input, "*.*")):
                if f.lower().endswith((".docx", ".doc")):
                    _process_auto(f, args.excel, args.row, args.out, args.dry_run)
        else:
            if not args.doc:
                raise SystemExit("Укажите --doc или --input")
            _process_auto(args.doc, args.excel, args.row, args.out, args.dry_run)
        return

    run_single(args)

if __name__ == "__main__":
    main()
