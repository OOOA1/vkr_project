# filldoc/detector.py
# -*- coding: utf-8 -*-
from dataclasses import dataclass
from typing import Tuple, Dict, List
import re
from docx.document import Document as _Document
from .templates import TEMPLATES, TemplateSpec
from .utils import normalize

@dataclass
class DetectResult:
    template_id: str | None
    score: int
    breakdown: List[str]

def _first_page_text(doc: _Document) -> Tuple[str, Dict[str,int]]:
    texts: List[str] = []
    tables = 0
    used = 0
    LIMIT = 2000

    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            texts.append(t)
            used += len(t)
            if used > LIMIT:
                break

    for tbl in doc.tables:
        tables += 1
        if tables > 3:
            break
        for row in tbl.rows:
            cell_txt = " | ".join([c.text.strip() for c in row.cells if c.text])
            if cell_txt:
                texts.append(cell_txt)
                used += len(cell_txt)
                if used > LIMIT:
                    break

    raw = "\n".join(texts)
    low = normalize(raw).lower()
    feats = {
        "slashes": raw.count("/"),
        "underscores": raw.count("_"),
        "tables": tables,
    }
    return low, feats

def _score(text: str, feats: Dict[str,int], spec: TemplateSpec) -> Tuple[int, List[str]]:
    s = 0
    notes: List[str] = []
    w_req, w_opt, w_neg, w_layout = 5, 1, -6, 1

    def has(needle: str) -> bool:
        return normalize(needle).lower() in text

    for a in spec.detect.required:
        if has(a):
            s += w_req; notes.append(f"+{w_req} '{a}'")
        else:
            notes.append(f"0 '{a}' (required отсутствует)")

    for a in spec.detect.optional:
        if has(a):
            s += w_opt; notes.append(f"+{w_opt} '{a}'")

    for a in spec.detect.negative:
        if has(a):
            s += w_neg; notes.append(f"{w_neg} NEG '{a}'")

    # layout hints
    layout = spec.detect.layout or {}
    if layout.get("must_have_slash") and feats.get("slashes",0) > 0:
        s += w_layout; notes.append(f"+{w_layout} has '/'")
    if "min_tables" in layout and feats.get("tables",0) >= int(layout["min_tables"]):
        s += w_layout; notes.append(f"+{w_layout} tables≥{layout['min_tables']}")
    if "min_underscores" in layout and feats.get("underscores",0) >= int(layout["min_underscores"]):
        s += w_layout; notes.append(f"+{w_layout} underscores≥{layout['min_underscores']}")

    return s, notes

import re

def detect_template(doc: _Document) -> DetectResult:
    text, feats = _first_page_text(doc)     # text — единая строка (norm), feats — признаки
    lines, _ = _first_page_lines(doc)       # уже есть в файле; возьмём нормализованные строки
    norm_lines = [normalize(l).lower() for l in lines]

    best_id, best_sc, best_notes = None, -10**9, []
    for tid, spec in TEMPLATES.items():
        notes = []
        sc = 0

        # 0) жёсткие запреты — если встречается хоть один → мимо
        for bad in getattr(spec.detect, "negative_hard", []) or []:
            if normalize(bad).lower() in text:
                notes.append(f"-INF HARD_NEG '{bad}'")
                sc = -10**6
                break
        if sc == -10**6:
            if sc > best_sc:
                best_id, best_sc, best_notes = tid, sc, notes
            continue

        # 1) must_all — все обязаны встретиться (в любом месте текста)
        missing = [m for m in (spec.detect.must_all or []) if normalize(m).lower() not in text]
        if missing:
            for m in missing:
                notes.append(f"0 MUST '{m}' отсутствует")
            sc = -10**6
        if sc == -10**6:
            if sc > best_sc: best_id, best_sc, best_notes = tid, sc, notes
            continue

        # 2) any_of — для каждой группы хотя бы одна фраза должна встретиться
        for grp in getattr(spec.detect, "any_of", []) or []:
            if not any(normalize(p).lower() in text for p in grp):
                notes.append(f"0 ANY_OF {grp} — ни одной нет")
                sc = -10**6
                break
        if sc == -10**6:
            if sc > best_sc: best_id, best_sc, best_notes = tid, sc, notes
            continue

        # 3) must_exact_lines — точные совпадения с целой строкой
        for s in getattr(spec.detect, "must_exact_lines", []) or []:
            ns = normalize(s).lower()
            if ns not in norm_lines:
                notes.append(f"0 MUST_LINE '{s}' — нет точной строки")
                sc = -10**6
                break
        if sc == -10**6:
            if sc > best_sc: best_id, best_sc, best_notes = tid, sc, notes
            continue

        # 4) must_regex
        for rgx in getattr(spec.detect, "must_regex", []) or []:
            if not re.search(rgx, "\n".join(lines), flags=re.IGNORECASE | re.MULTILINE):
                notes.append(f"0 MUST_REGEX /{rgx}/ — нет совпадения")
                sc = -10**6
                break
        if sc == -10**6:
            if sc > best_sc: best_id, best_sc, best_notes = tid, sc, notes
            continue

        # 5) require_same_line — обе подстроки в ОДНОЙ строке
        for a, b in getattr(spec.detect, "require_same_line", []) or []:
            an, bn = normalize(a).lower(), normalize(b).lower()
            if not any((an in L and bn in L) for L in norm_lines):
                notes.append(f"0 SAME_LINE '{a}' + '{b}' — нет пары в строке")
                sc = -10**6
                break
        if sc == -10**6:
            if sc > best_sc: best_id, best_sc, best_notes = tid, sc, notes
            continue

        # 6) обычное скорирование required/optional + layout
        s2, notes2 = _score(text, feats, spec)
        sc += s2
        notes.extend(notes2)

        # 7) мягкие негативы — штраф, но не запрет
        for bad in getattr(spec.detect, "negative_soft", []) or []:
            if normalize(bad).lower() in text:
                sc -= 6
                notes.append(f"-6 SOFT_NEG '{bad}'")

        if sc > best_sc:
            best_id, best_sc, best_notes = tid, sc, notes

    if best_id is None:
        return DetectResult(None, -10**9, best_notes)

    thr = TEMPLATES[best_id].detect.threshold
    if best_sc < thr:
        return DetectResult(None, best_sc, best_notes)

    return DetectResult(best_id, best_sc, best_notes)

def _first_page_lines(doc) -> Tuple[List[str], Dict[str,int]]:
    """Вернуть список строк первой страницы (≈ первые абзацы + до 3 таблиц)
    + признаки: кол-во '/', '_' и таблиц."""
    lines: List[str] = []
    tables = 0
    used = 0
    LIMIT = 2000

    # абзацы
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(t)
            used += len(t)
            if used > LIMIT:
                break

    # немного таблиц (по строкам)
    for tbl in doc.tables:
        tables += 1
        if tables > 3:
            break
        for row in tbl.rows:
            t = " | ".join([(c.text or "").strip() for c in row.cells])
            t = t.strip()
            if t:
                lines.append(t)
                used += len(t)
                if used > LIMIT:
                    break

    raw = "\n".join(lines)
    feats = {
        "slashes": raw.count("/"),
        "underscores": raw.count("_"),
        "tables": tables,
    }
    return lines, feats

def suggest_anchors(doc) -> Dict[str, List[str]]:
    """Подсказать якоря для detect.required/optional."""
    lines, feats = _first_page_lines(doc)
    # нормализованные для совпадений
    norm = [normalize(l) for l in lines]

    top_lines = [normalize(l) for l in lines[:10]]

    with_colon = []
    seen = set()
    for l in norm:
        if ":" in l and l not in seen:
            with_colon.append(l)
            seen.add(l)
        if len(with_colon) >= 20:
            break

    uppercase = [l for l in norm if l == l.upper() and len(l) >= 4][:10]

    keywords = []
    keys = ("студент", "руковод", "кафедр", "группа", "курс",
            "направлен", "специальн", "тема", "практик", "вкр",
            "заявлен", "согласован", "эбс", "каникул", "лист")
    for l in norm:
        if any(k in l for k in keys):
            keywords.append(l)
    keywords = list(dict.fromkeys(keywords))[:20]

    return {
        "top_lines": top_lines,
        "with_colon": with_colon,
        "uppercase": uppercase,
        "keywords": keywords,
    }