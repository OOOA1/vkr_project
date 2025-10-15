# filldoc/cli.py
# -*- coding: utf-8 -*-
import argparse, os, glob, re
from docx import Document

from .utils import convert_doc_to_docx
from .detector import detect_template, suggest_anchors
from .templates import TEMPLATES


# ---------------------- вспомогалки ---------------------- #

def render_mask(mask: str, row: dict) -> str:
    def _repl(m):
        key = m.group(1).strip()
        val = row.get(key)
        return "" if val is None else str(val)
    return re.sub(r"\{\{\s*([^}]+?)\s*\}\}", _repl, mask)


def _safe_title(template_id: str) -> str:
    return TEMPLATES.get(template_id, template_id)


# ---------------------- совместимость с excel.py ---------------------- #

class _TemplateSpec:
    def __init__(self, template_id: str, mapping_rows: list[dict], settings: dict):
        self.id = template_id
        self.doc_type = template_id
        self.mapping = mapping_rows or []
        self.settings = settings or {}

class _MasterCompat:
    """
    Унифицированная обёртка над любым форматом, который вернёт excel.load_master_data.
    Всегда предоставляет:
      - .data_rows: list[dict]
      - .for_template(doc_type) -> _TemplateSpec
    """
    def __init__(self, data_rows: list[dict], mapping_rows: list[dict], settings: dict):
        self.data_rows = data_rows or []
        self._mapping_rows = mapping_rows or []
        self._settings = settings or {}

    def for_template(self, doc_type: str) -> _TemplateSpec | None:
        mm = [r for r in self._mapping_rows if str(r.get("doc_type", "")).strip() == str(doc_type)]
        if not mm:
            return None
        return _TemplateSpec(doc_type, mm, self._settings)


def _load_master_any(excel_path: str) -> _MasterCompat:
    """
    Пытается использовать твой excel.load_master_data(...) как есть.
    Если он не возвращает объект с .for_template/.data_rows — подчитывает Excel самостоятельно.
    """
    # 1) Пробуем твой excel.load_master_data
    try:
        from .excel import load_master_data  # type: ignore
        master = load_master_data(excel_path)
        # уже подходящий объект?
        if hasattr(master, "for_template") and hasattr(master, "data_rows"):
            return master  # type: ignore
        # может быть список спецификаций + отдельные функции?
        # если это список или что-то "сырая структура" — пойдём в наш fallback
    except Exception:
        # Падает — просто пойдём в наш fallback
        pass

    # 2) Fallback: читаем Excel напрямую (settings/data/mapping)
    try:
        from openpyxl import load_workbook
    except Exception as e:
        raise SystemExit("Нужен openpyxl для чтения Excel. Установи пакет: pip install openpyxl") from e

    if not os.path.exists(excel_path):
        raise SystemExit(f"Excel не найден: {excel_path}")

    wb = load_workbook(excel_path, data_only=True)

    # settings: key|value
    settings = {}
    if "settings" in wb.sheetnames:
        ws = wb["settings"]
        # ожидаем 2 колонки: key, value
        # игнорируем пустые строки
        headers = [str(c.value).strip().lower() if c.value is not None else "" for c in next(ws.iter_rows(max_row=1))]
        # допускаем порядки/названия в любом регистре
        try:
            key_idx = headers.index("key")
            val_idx = headers.index("value")
            for row in ws.iter_rows(min_row=2):
                k = row[key_idx].value
                v = row[val_idx].value
                if k is not None:
                    settings[str(k).strip()] = v if v is not None else ""
        except ValueError:
            # нет нормальных заголовков — тихо игнорируем, будет пусто
            pass

    # data: любая шапка, обязателен RecordID
    data_rows: list[dict] = []
    if "data" in wb.sheetnames:
        ws = wb["data"]
        headers = [str(c.value).strip() if c.value is not None else "" for c in next(ws.iter_rows(max_row=1))]
        for row in ws.iter_rows(min_row=2):
            d = {}
            for i, cell in enumerate(row):
                col = headers[i]
                if col:
                    d[col] = cell.value
            # фильтруем полностью пустые
            if any(v not in (None, "", " ") for v in d.values()):
                data_rows.append(d)

    # mapping: ожидаем колонки doc_type, selector_type, selector, direction, key, fmt, when
    mapping_rows: list[dict] = []
    if "mapping" in wb.sheetnames:
        ws = wb["mapping"]
        headers = [str(c.value).strip() if c.value is not None else "" for c in next(ws.iter_rows(max_row=1))]
        # нормализуем известные синонимы
        normalize = {
            "doctype": "doc_type", "тип_документа": "doc_type",
            "type": "selector_type", "тип_селектора": "selector_type",
            "селектор": "selector",
            "dir": "direction", "направление": "direction",
            "field": "key", "поле": "key",
            "format": "fmt", "формат": "fmt",
            "cond": "when", "условие": "when",
        }
        headers_std = [normalize.get(h.lower(), h) for h in headers]
        for row in ws.iter_rows(min_row=2):
            rec = {}
            for i, cell in enumerate(row):
                col = headers_std[i]
                if col:
                    rec[col] = cell.value if cell.value is not None else ""
            # минимальная валидация
            if str(rec.get("doc_type", "")).strip() and str(rec.get("selector_type", "")).strip() and str(rec.get("selector", "")).strip():
                mapping_rows.append(rec)

    return _MasterCompat(data_rows=data_rows, mapping_rows=mapping_rows, settings=settings)


# ---------------------- режимы ---------------------- #

def _process_detect(path: str):
    try:
        src = convert_doc_to_docx(path)
    except Exception as e:
        print(f"[CONVERT-ERROR] {os.path.basename(path)}: {e}")
        return
    base = os.path.basename(path)
    doc = Document(src)
    det = detect_template(doc)
    if not det.template_id:
        print(f"[DETECT] '{base}' → unknown")
        return
    title = _safe_title(det.template_id)
    print(f"[DETECT] '{base}' → {det.template_id} ({title}), score={det.score}")
    for w in det.why:
        print(f"  - {w}")


def _process_suggest(path: str):
    try:
        src = convert_doc_to_docx(path)
    except Exception as e:
        print(f"[CONVERT-ERROR] {os.path.basename(path)}: {e}")
        return
    base = os.path.basename(path)
    doc = Document(src)
    det = detect_template(doc)
    if not det.template_id:
        print(f"[SUGGEST] '{base}' → unknown (не удалось распознать шаблон)")
        return
    anchors = suggest_anchors(doc)
    print(f"[SUGGEST] '{base}' ({det.template_id}) возможные подписи/якоря:")
    for a in anchors:
        print(" ", a)


def _process_auto(path: str, excel: str, row: int | None, outdir: str, dry_run: bool):
    # генератор
    try:
        from .generator import apply_mapping_to_doc
    except Exception as e:
        raise SystemExit("В filldoc/generator.py должна быть функция apply_mapping_to_doc(doc, data_row, mapping_rows, settings, log=None, dry_run=False).") from e

    # конвертация/открытие
    try:
        src = convert_doc_to_docx(path)
    except Exception as e:
        print(f"[CONVERT-ERROR] {os.path.basename(path)}: {e}")
        return
    base = os.path.basename(path)
    doc = Document(src)
    det = detect_template(doc)
    if not det.template_id:
        print(f"[AUTO] '{base}' → пропуск (unknown)")
        return

    # Excel → совместимый мастер
    master = _load_master_any(excel)
    spec = master.for_template(det.template_id)
    if not spec or not spec.mapping:
        print(f"[{det.template_id}] '{base}' → нет правил в mapping (doc_type={det.template_id})")
        return

    # данные
    rows = master.data_rows
    if row is not None:
        idx = int(row)
        if idx < 1 or idx > len(rows):
            raise SystemExit(f"--row={idx} вне диапазона 1..{len(rows)}")
        rows = [rows[idx - 1]]

    os.makedirs(outdir, exist_ok=True)

    # Маска имени файла
    mask = spec.settings.get("маска_имени_файла") \
           or spec.settings.get("output_mask") \
           or "{{doc_type}}_{{ФИО}}.docx"

    for ridx, r in enumerate(rows, start=1):
        # Имя файла
        filename_vars = dict(r)
        filename_vars["doc_type"] = det.template_id
        out_name = render_mask(mask, filename_vars)
        if not out_name.lower().endswith(".docx"):
            out_name += ".docx"
        out_path = os.path.join(outdir, out_name)

        # Применение правил
        log: list[str] = []
        d = Document(src)

        # ВАЖНО: передаём log=log, чтобы dry-run печатал замену
        apply_mapping_to_doc(d, r, spec.mapping, spec.settings, log=log, dry_run=dry_run)

        # Превью/сохранение
        print(f"[{spec.id}] Предпросмотр '{out_name}':" if dry_run else f"[{spec.id}] Сохранение '{out_name}':")
        for line in log:
            print("  -", line)
        if not dry_run:
            d.save(out_path)
            print(f"  -> Сохранено: {out_path}")


def main():
    p = argparse.ArgumentParser()
    p.add_argument("--detect", action="store_true", help="распознать шаблон по 1-й странице")
    p.add_argument("--suggest", action="store_true", help="подсказать якоря/подписи на 1-й странице")
    p.add_argument("--auto", action="store_true", help="автозаполнение по Excel master")
    p.add_argument("--doc", help="один документ (.docx/.doc)")
    p.add_argument("--input", help="папка с документами")
    p.add_argument("--excel", help="путь к master.xlsx", default="master.xlsx")
    p.add_argument("--row", type=int, help="номер строки из листа data (1-based)")
    p.add_argument("--out", help="папка для вывода", default="./output")
    p.add_argument("--dry-run", action="store_true", help="только лог, без сохранения изменений")
    args = p.parse_args()

    if args.detect:
        if args.input:
            for f in glob.glob(os.path.join(args.input, "*.*")):
                if f.lower().endswith((".docx", ".doc")):
                    _process_detect(f)
        else:
            if not args.doc:
                raise SystemExit("Укажите --doc или --input")
            _process_detect(args.doc)
        return

    if args.suggest:
        if args.input:
            for f in glob.glob(os.path.join(args.input, "*.*")):
                if f.lower().endswith((".docx", ".doc")):
                    _process_suggest(f)
        else:
            if not args.doc:
                raise SystemExit("Укажите --doc или --input")
            _process_suggest(args.doc)
        return

    if args.auto:
        if args.input:
            for f in glob.glob(os.path.join(args.input, "*.*")):
                if f.lower().endswith((".docx", ".doc")):
                    _process_auto(f, args.excel, args.row, args.out, args.dry_run)
        else:
            if not args.doc:
                raise SystemExit("Укажите --doc или --input")
            _process_auto(args.doc, args.excel, args.row, args.dry_run)
        return

    # Режим одиночного, если реализован
    try:
        from .cli_single import run_single
    except Exception:
        raise SystemExit("Укажите режим: --detect | --suggest | --auto")
    run_single(args)


if __name__ == "__main__":
    main()
