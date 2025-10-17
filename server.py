# server.py
# Зависимости:
#   pip install fastapi "uvicorn[standard]" python-multipart pandas openpyxl docxtpl requests
# (docxtpl тянет python-docx, используется для генерации DOCX-инструкции)

import io
import re
import csv
import zipfile
from pathlib import Path
from typing import Optional, Dict, Tuple, List

import pandas as pd
import requests
from fastapi import FastAPI, File, Form, UploadFile, HTTPException
from fastapi.responses import (
    HTMLResponse,
    JSONResponse,
    StreamingResponse,
    PlainTextResponse,
    FileResponse,
)
from docxtpl import DocxTemplate

# для генерации DOCX-инструкции (ставится вместе с docxtpl)
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from templates_config import TEMPLATES  # твои захардкоженные шаблоны

app = FastAPI(title="VRK Docs — DOCX → ZIP", version="3.6.0")

# === Пути базы ===
BASE_DIR = Path(__file__).resolve().parent

# === Настройки выдачи Excel-шаблона ===
TEMPLATE_DOWNLOAD_NAME = "main_example.xlsx"
CANDIDATE_TEMPLATE_PATHS: List[Path] = [
    BASE_DIR / "main_example.xlsx",
    BASE_DIR / "main.xlsx",
    BASE_DIR / "main — копия.xlsx",
]

# === Настройки выдачи Инструкции ===
INSTRUCTION_DOWNLOAD_NAME = "instruction.docx"
INSTRUCTION_CANDIDATES: List[Path] = [
    BASE_DIR / "instruction.docx",
    BASE_DIR / "инструкция.docx",
    BASE_DIR / "instruction.doc",   # если вдруг положите .doc
    BASE_DIR / "instruction.dock",  # опечатка — тоже поддержим
]

# ============= Красивый UI (без внешних зависимостей) =============
INDEX_HTML = """
<!doctype html><meta charset="utf-8">
<title>VRK Docs — DOCX → ZIP</title>
<meta name="viewport" content="width=device-width, initial-scale=1">
<style>
  :root{
    --bg:#0b1220; --card:#0f172a; --muted:#94a3b8; --text:#e5e7eb;
    --accent:#60a5fa; --accent-2:#22d3ee; --accent-3:#a78bfa;
    --success:#10b981; --warn:#f59e0b; --danger:#ef4444;
    --border:#1f2a44; --shadow:0 20px 40px rgba(0,0,0,.35);
    --radius:18px;
  }
  @media (prefers-color-scheme: light) {
    :root{
      --bg:#f6f7fb; --card:#ffffff; --text:#0f172a; --muted:#64748b;
      --border:#e5e7eb; --shadow:0 18px 40px rgba(2,6,23,.08);
    }
  }
  *{box-sizing:border-box}
  html,body{height:100%}
  body{
    margin:0; font-family: ui-sans-serif,system-ui,-apple-system,Segoe UI,Roboto,Arial;
    background:
      radial-gradient(1200px 400px at 10% -10%, rgba(96,165,250,.12), transparent 60%),
      radial-gradient(900px 300px at 100% 20%, rgba(167,139,250,.12), transparent 60%),
      var(--bg);
    color:var(--text);
  }
  .wrap{max-width:1100px; margin:48px auto; padding:0 20px}
  .hero{display:flex; align-items:flex-start; gap:18px; margin-bottom:24px;}
  .logo{
    width:48px; height:48px; border-radius:50%;
    background:linear-gradient(135deg,var(--accent),var(--accent-3));
    display:grid; place-items:center; color:white; font-weight:800;
    box-shadow:var(--shadow);
  }
  .title{font-size:28px; font-weight:800; margin:0}
  .subtitle{margin:4px 0 0; color:var(--muted)}
  .card{
    background:var(--card); border:1px solid var(--border);
    border-radius:var(--radius); padding:22px; box-shadow:var(--shadow);
  }
  .grid{display:grid; grid-template-columns:1fr 1fr; gap:18px}
  @media (max-width:900px){ .grid{grid-template-columns:1fr} }
  label{font-weight:700; font-size:14px; margin-bottom:8px; display:block}
  input[type=file], input[type=url], input[type=number]{
    width:100%; padding:12px 14px; border-radius:12px; outline:0;
    border:1px solid var(--border); background:transparent; color:var(--text);
  }
  input::placeholder{color:var(--muted)}
  .drop{border:1.6px dashed var(--border); border-radius:14px; padding:18px;
    display:flex; align-items:center; gap:12px; transition:.2s;}
  .drop.drag{border-color:var(--accent)}
  .chip{display:inline-flex; align-items:center; gap:6px; padding:6px 10px; border-radius:999px;
    font-size:12px; border:1px solid var(--border); color:var(--muted); background:transparent}
  .row{display:flex; gap:12px; flex-wrap:wrap; align-items:center}
  .btn{
    appearance:none; border:0; cursor:pointer; color:#0b1220; font-weight:800;
    padding:12px 18px; border-radius:12px; transition:transform .05s ease;
    background:linear-gradient(135deg,var(--accent),var(--accent-2));
  }
  .btn.secondary{background:transparent; color:var(--text); border:1px solid var(--border)}
  .btn:active{transform:translateY(1px)}
  .btn:disabled{opacity:.6; cursor:not-allowed}
  .badges{display:flex; gap:8px; flex-wrap:wrap}
  .badge{font-size:12px; padding:4px 10px; border-radius:999px; border:1px solid var(--border); color:var(--muted)}
  .badge.ok{color:var(--success); border-color:rgba(16,185,129,.35)}
  .badge.warn{color:var(--warn); border-color:rgba(245,158,11,.35)}
  .badge.err{color:var(--danger); border-color:rgba(239,68,68,.35)}
  .preview{margin-top:16px; overflow:auto; border:1px solid var(--border); border-radius:12px}
  table{border-collapse:collapse; width:100%; font-size:13px}
  th,td{border-bottom:1px solid var(--border); padding:10px 12px; text-align:left}
  th{color:var(--muted); font-weight:700}
  .foot{margin-top:18px; color:var(--muted); font-size:12px}
  .toast{margin-top:12px; padding:12px 14px; border-radius:12px; border:1px solid var(--border)}
  .toast.ok{border-color:rgba(16,185,129,.35); color:var(--success)}
  .toast.err{border-color:rgba(239,68,68,.35); color:var(--danger)}
  .spinner{width:16px;height:16px;border-radius:999px;border:3px solid rgba(255,255,255,.25);border-top-color:#fff;animation:spin .8s linear infinite}
  @keyframes spin{to{transform:rotate(360deg)}}
  .actions{display:flex; gap:10px; align-items:center; justify-content:flex-start; flex-wrap:wrap; margin-top:8px}
</style>

<div class="wrap">
  <div class="hero">
    <div class="logo">V</div>
    <div>
      <h1 class="title">VRK Docs — генерация DOCX → ZIP</h1>
      <p class="subtitle">Загрузите Excel/CSV <b>или</b> укажите Google Sheet. «Широкая» таблица: <b>строка 1 — заголовки</b>, <b>строка 2 — значения</b>. Если обнаружим «шпаргалку» вида <code>output_mask</code>, переключимся в режим «ключ → значение».</p>
    </div>
  </div>

  <div class="card">
    <form id="f" action="/generate" method="post" enctype="multipart/form-data">
      <div class="grid">
        <div>
          <label>Excel (.xlsx) или CSV</label>
          <div class="drop" id="drop">
            <svg width="22" height="22" viewBox="0 0 24 24" fill="none"><path d="M12 16V4m0 12 4-4m-4 4-4-4M4 16v2a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2v-2" stroke="currentColor" stroke-width="1.7" stroke-linecap="round" stroke-linejoin="round"/></svg>
            <div>
              <input type="file" name="table_file" id="table_file" accept=".xlsx,.csv" style="display:none">
              <button class="btn secondary" id="chooseFile" type="button">Выбрать файл</button>
              <div class="hint" id="fileHint">Файл не выбран</div>
            </div>
          </div>
        </div>
        <div>
          <label>Google Sheet (опционально)</label>
          <input type="url" name="gsheet_url" id="gsheet_url" placeholder="https://docs.google.com/spreadsheets/d/.../edit#gid=0">
          <div class="hint">Доступ к таблице: <b>Anyone with the link</b>. Если введёте ссылку — файл очищаем автоматически.</div>
        </div>
      </div>

      <div class="grid" style="margin-top:16px">
        <div>
          <label>Строка заголовков (для «широкой» таблицы)</label>
          <input type="number" name="header_row" id="header_row" value="1" min="1">
        </div>
        <div>
          <label>Статус источника</label>
          <div class="row">
            <span class="chip" id="chipSource">источник не выбран</span>
            <span class="chip" id="chipMode">режим: —</span>
          </div>
        </div>
      </div>

      <div class="actions">
        <button id="btnInspect" type="button" class="btn">Проверить</button>
        <button id="btnGen" type="submit" class="btn">Сгенерировать ZIP</button>
        <button id="btnTpl" class="btn secondary" type="button">Скачать шаблон (main_example.xlsx)</button>
        <button id="btnInstr" class="btn secondary" type="button">Скачать инструкцию (DOCX)</button>
        <span id="spin" style="display:none" class="spinner"></span>
      </div>
    </form>

    <div id="rep" class="foot"></div>
  </div>
</div>

<script>
  const $ = (s)=>document.querySelector(s);
  const f = $('#f'), rep = $('#rep'), spin = $('#spin');
  const file = $('#table_file'), gs = $('#gsheet_url'), headerRow=$('#header_row');
  const chooseBtn=$('#chooseFile'), fileHint=$('#fileHint'), drop=$('#drop');
  const chipSource=$('#chipSource'), chipMode=$('#chipMode');
  const btnI=$('#btnInspect'), btnG=$('#btnGen'), btnTpl=$('#btnTpl'), btnInstr=$('#btnInstr');

  // Remember header row
  headerRow.value = localStorage.getItem('hdrRow') || '1';
  headerRow.addEventListener('change', ()=>localStorage.setItem('hdrRow', headerRow.value));

  function setSourceChip(){
    const hasFile = file.files && file.files.length>0;
    const hasGS = !!gs.value.trim();
    if(hasGS){
      chipSource.textContent = 'источник: Google Sheet';
    }else if(hasFile){
      chipSource.textContent = 'источник: файл ' + (file.files[0]?.name || '');
    }else{
      chipSource.textContent = 'источник не выбран';
    }
  }

  function needSource(){
    const hasFile = file.files && file.files.length>0;
    const hasGS = !!gs.value.trim();
    if(!hasFile && !hasGS){
      toast('Укажите Google Sheet ИЛИ выберите файл', 'err');
      return false;
    }
    return true;
  }

  function toast(text, kind){
    rep.innerHTML = '<div class="toast '+(kind||'')+'">'+text+'</div>' + rep.innerHTML;
  }

  // File select & DnD
  chooseBtn.addEventListener('click', ()=> file.click());
  file.addEventListener('change', ()=>{
    if(file.files.length){ gs.value=''; fileHint.textContent = 'Выбран: '+file.files[0].name; }
    else { fileHint.textContent = 'Файл не выбран'; }
    setSourceChip();
  });
  gs.addEventListener('input', ()=>{ if(gs.value.trim()) { file.value=''; fileHint.textContent='Файл не выбран'; } setSourceChip(); });

  ;['dragenter','dragover'].forEach(ev=>drop.addEventListener(ev, e=>{e.preventDefault(); drop.classList.add('drag');}));
  ;['dragleave','drop'].forEach(ev=>drop.addEventListener(ev, e=>{e.preventDefault(); drop.classList.remove('drag');}));
  drop.addEventListener('drop', (e)=>{
    const dt=e.dataTransfer; if(!dt||!dt.files||!dt.files.length) return;
    const f=dt.files[0];
    if(!/\.(xlsx|csv)$/i.test(f.name)){ toast('Ожидается .xlsx или .csv', 'err'); return; }
    file.files = dt.files; gs.value=''; fileHint.textContent='Выбран: '+f.name; setSourceChip();
  });

  function kvPairs(pairs){
    if(!pairs||!pairs.length) return '';
    const rows = pairs.map(([k,v])=>`<tr><td>${k}</td><td>${v}</td></tr>`).join('');
    return `<div class="preview"><table><thead><tr><th>Ключ</th><th>Значение</th></tr></thead><tbody>${rows}</tbody></table></div>`;
  }
  function previewWide(obj){
    if(!obj) return '';
    const rows = Object.entries(obj).map(([k,v])=>`<tr><th>${k}</th><td>${v}</td></tr>`).join('');
    return `<div class="preview"><table><tbody>${rows}</tbody></table></div>`;
  }

  async function inspect(){
    if(!needSource()) return;
    rep.innerHTML=''; spin.style.display='inline';
    const fd=new FormData(f);
    const r=await fetch('/inspect',{method:'POST',body:fd});
    const d=await r.json(); spin.style.display='none';
    if(!r.ok){ toast(d.detail||'Ошибка', 'err'); return; }

    chipMode.textContent = 'режим: '+(d.meta?.mode || '—');

    const meta = d.meta ? `<div class="badges">
      <span class="badge">источник: ${d.meta.source}</span>
      ${d.meta.mode==='wide' ? `<span class="badge">header: ${(d.meta.header_row??0)+1}</span>` : ''}
      ${d.meta.gid!==undefined ? `<span class="badge">gid: ${d.meta.gid}</span>` : ''}
    </div>` : '';

    const miss = d.missing?.length
      ? '<div class="toast warn">Отсутствуют ключевые поля: '+d.missing.join(', ')+'</div>'
      : '<div class="toast ok">Ключевые поля найдены</div>';

    const cols = (d.columns?.length)
      ? '<div class="foot"><b>Колонки ('+d.columns.length+'):</b> '+d.columns.join(', ')+'</div>' : '';

    const body = d.meta?.mode==='wide' ? previewWide(d.preview) : kvPairs(d.preview_pairs);
    rep.innerHTML = meta + miss + body + cols;
  }

  btnTpl.addEventListener('click', async ()=>{
    try{
      const r = await fetch('/template', { method:'GET', cache:'no-store' });
      if(!r.ok){ const t=await r.text(); toast('Не удалось скачать шаблон: '+t,'err'); return; }
      const b = await r.blob(); const url=URL.createObjectURL(b);
      const a = document.createElement('a'); a.href=url; a.download='main_example.xlsx';
      document.body.appendChild(a); a.click(); a.remove(); URL.revokeObjectURL(url);
      toast('Шаблон скачан', 'ok');
    }catch(err){ toast('Ошибка скачивания: '+err,'err'); }
  });

  btnInstr.addEventListener('click', async ()=>{
    try{
      const r = await fetch('/instruction', { method:'GET', cache:'no-store' });
      if(!r.ok){ const t=await r.text(); toast('Не удалось скачать инструкцию: '+t,'err'); return; }
      const b = await r.blob(); const url=URL.createObjectURL(b);
      const a = document.createElement('a'); a.href=url; a.download='instruction.docx';
      document.body.appendChild(a); a.click(); a.remove(); URL.revokeObjectURL(url);
      toast('Инструкция скачана', 'ok');
    }catch(err){ toast('Ошибка скачивания: '+err,'err'); }
  });

  btnI.addEventListener('click', inspect);

  f.addEventListener('submit', async (e)=>{
    e.preventDefault();
    const hasFile = file.files && file.files.length>0;
    const hasGS = !!gs.value.trim();
    if(!hasFile && !hasGS){ toast('Укажите Google Sheet ИЛИ выберите файл','err'); return; }
    rep.innerHTML=''; btnI.disabled=btnG.disabled=true; spin.style.display='inline';
    try{
      const fd=new FormData(f); const r=await fetch('/generate',{method:'POST',body:fd}); const b=await r.blob();
      if(!r.ok){ rep.innerHTML=''; toast(await b.text(),'err'); return; }
      const url=URL.createObjectURL(b); const a=document.createElement('a'); a.href=url; a.download='generated_docs.zip';
      document.body.appendChild(a); a.click(); a.remove(); URL.revokeObjectURL(url);
      toast('ZIP готов и скачан','ok');
    } catch(err){
      toast('Сетевая ошибка: '+err,'err');
    } finally {
      btnI.disabled=btnG.disabled=false; spin.style.display='none';
    }
  });

  // init
  setSourceChip();
</script>
"""

# ============= Бизнес-логика =============
INVALID_FS = r'[<>:"/\\|?*]'

def safe(v): return "" if (v is None or pd.isna(v)) else str(v).strip()

class SafeMap(dict):
    def __missing__(self, key): return ""

def slugify(name: str) -> str:
    return re.sub(INVALID_FS, "_", name).rstrip(" .") or "file"

def _norm(s: str) -> str:
    return re.sub(r"\s+", "", str(s)).replace("\ufeff","").replace("\xa0","").replace("ё","е").lower()

def expected_headers() -> set:
    exp = {"фио","группа"}
    for tpl in TEMPLATES:
        exp |= {_norm(v) for v in tpl["fields"].values()}
        exp |= {_norm(m) for m in re.findall(r"\{([^}]+)\}", tpl["out"])}
    return exp

def score_columns(cols) -> int:
    exp = expected_headers()
    return sum(1 for c in cols if _norm(c) in exp)

def read_wide_try(file_bytes: bytes, is_xlsx: bool, header_row: int) -> Tuple[pd.DataFrame, Dict]:
    if is_xlsx:
        df = pd.read_excel(io.BytesIO(file_bytes), sheet_name=0, header=max(header_row-1,0))
        return df, {"source":"xlsx", "mode":"wide", "header_row": header_row-1}
    else:
        sample = file_bytes[:2048].decode("utf-8", errors="ignore")
        try: sep = csv.Sniffer().sniff(sample).delimiter
        except Exception: sep = ","
        df = pd.read_csv(io.BytesIO(file_bytes), sep=sep, header=max(header_row-1,0))
        return df, {"source":"csv", "mode":"wide", "header_row": header_row-1}

def read_kv_from_raw(file_bytes: bytes, is_xlsx: bool, key_row: int = 1, val_row: int = 2) -> Tuple[Dict[str,str], Dict]:
    if is_xlsx:
        df = pd.read_excel(io.BytesIO(file_bytes), sheet_name=0, header=None)
    else:
        df = pd.read_csv(io.BytesIO(file_bytes), header=None)
    keys = [safe(x).replace("\ufeff","").replace("\xa0"," ") for x in df.iloc[key_row-1].tolist()]
    vals = [safe(x).replace("\ufeff","").replace("\xa0"," ") for x in df.iloc[val_row-1].tolist()]
    kv = {k: v for k, v in zip(keys, vals) if k}
    return kv, {"source":"xlsx" if is_xlsx else "csv", "mode":"kv", "key_row":key_row-1, "val_row":val_row-1}

def extract_record_from_upload(file: UploadFile, header_row: int) -> Tuple[Dict[str,str], Dict, Optional[list]]:
    data = file.file.read()
    name = (file.filename or "").lower()
    is_xlsx = name.endswith(".xlsx")
    if not (is_xlsx or name.endswith(".csv")):
        raise HTTPException(400, "Поддерживаются только .xlsx или .csv")

    df_wide, meta = read_wide_try(data, is_xlsx, header_row)
    if not df_wide.empty:
        cols = [str(c) for c in df_wide.columns]
        sc = score_columns(cols)
        if sc >= 3:
            row = pick_first_nonempty_row(df_wide)
            row_dict = {str(k): safe(v) for k, v in row.items()}
            meta.update({"mode":"wide", "score": sc})
            return row_dict, meta, cols

    kv, meta_kv = read_kv_from_raw(data, is_xlsx, 1, 2)
    meta_kv.setdefault("score", 0)
    return kv, meta_kv, None

def extract_record_from_gsheet(url: str, header_row: int) -> Tuple[Dict[str,str], Dict, Optional[list]]:
    m = re.search(r"/spreadsheets/d/([a-zA-Z0-9-_]+)", url or "")
    if not m: raise HTTPException(400, "Не удалось извлечь spreadsheetId из URL")
    spreadsheet_id = m.group(1)
    gid_match = re.search(r"[#&?]gid=([0-9]+)", url)
    gid = int(gid_match.group(1)) if gid_match else 0
    export = f"https://docs.google.com/spreadsheets/d/{spreadsheet_id}/export?format=csv&gid={gid}"
    resp = requests.get(export, timeout=30)
    if resp.status_code != 200:
        raise HTTPException(400, f"Google Sheets недоступен (HTTP {resp.status_code})")
    upl = UploadFile(filename="gs.csv", file=io.BytesIO(resp.content))
    rec, meta, cols = extract_record_from_upload(upl, header_row)
    meta.update({"source":"gsheet", "gid": gid})
    return rec, meta, cols

def pick_first_nonempty_row(df: pd.DataFrame) -> pd.Series:
    df = df.fillna("")
    for _, row in df.iterrows():
        if any(safe(v) for v in row.values):
            return row
    raise HTTPException(400, "Не найдена ни одна непустая строка с данными")

# -------- шаблон Excel --------
@app.get("/template")
def download_template():
    """
    Отдаём шаблон main_example.xlsx из корня проекта.
    Если не найден — генерируем минимальный шаблон из полей templates_config.py.
    Во всех случаях имя скачивания = main_example.xlsx. Кэш выключен.
    """
    for p in CANDIDATE_TEMPLATE_PATHS:
        if p.exists():
            return FileResponse(
                str(p),
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                filename=TEMPLATE_DOWNLOAD_NAME,
                headers={"Cache-Control": "no-store, no-cache, must-revalidate"},
            )
    # fallback
    headers_list: List[str] = []
    seen = set()
    for tpl in TEMPLATES:
        for col in tpl["fields"].values():
            if col not in seen:
                seen.add(col)
                headers_list.append(col)
    df = pd.DataFrame([[""] * len(headers_list)], columns=headers_list)
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as w:
        df.to_excel(w, sheet_name="Sheet1", index=False)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={
            "Content-Disposition": f'attachment; filename="{TEMPLATE_DOWNLOAD_NAME}"',
            "Cache-Control": "no-store, no-cache, must-revalidate",
        },
    )

# -------- инструкция DOCX --------
def _build_instruction_docx_bytes() -> bytes:
    """Генерация дефолтной инструкции (если нет готового файла)."""
    doc = Document()
    # стиль
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    doc.add_heading("Инструкция по заполнению Excel-таблицы main_example.xlsx", 0)

    p = doc.add_paragraph(
        "Таблица содержит один лист. В первой строке расположены названия полей, "
        "во второй строке — значения для одного студента. На основании этих значений "
        "формируются все документы Word из набора шаблонов."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_heading("Общие правила", level=1)
    rules = [
        "Формат даты: ДД.ММ.ГГГГ (например, 17.10.2025).",
        "Телефон в формате +7 999 123-45-67 или 8 999 123-45-67.",
        "E-mail: user@example.org.",
        "ФИО в именительном падеже: Иванов Иван Иванович.",
        "Поле «Курс» — целое число (1, 2, 3, 4...).",
        "ИНН организации — 10 или 12 цифр, без пробелов.",
        "Адреса указывайте полностью, как в официальных документах.",
        "Если колонка продублирована (например, «АдресОрганизации» и «АдрессОрганизации»), укажите одно и то же значение в обеих.",
    ]
    for r in rules:
        doc.add_paragraph(r, style=None).paragraph_format.left_indent = Pt(14)

    doc.add_heading("Список основных полей", level=1)
    fields = [
        ("ФИО", "Полное имя студента (И.П.). Пример: Иванов Иван Иванович."),
        ("Группа", "Учебная группа. Пример: Изу-101."),
        ("Курс", "Номер курса. Пример: 3."),
        ("ТипПрактики / ВидПрактика", "Например: производственная, преддипломная."),
        ("НачалоПрактики / КонецПрактики", "Даты в формате ДД.ММ.ГГГГ."),
        ("БазаПрактики", "Наименование организации. Пример: ООО «Ромашка»."),
        ("АдресОрганизации / АдрессОрганизации", "Почтовый адрес организации (одно и то же значение в обеих колонках)."),
        ("ЮрАдресПрофОрг", "Юридический адрес учебного подразделения."),
        ("ОргИНН", "ИНН организации."),
        ("РукПрофОрг / РукВУЗФИО / РукВУЗ", "ФИО/должности руководителей от организации и вуза."),
        ("Кафедра / КафедраРП", "Название кафедры."),
        ("Научный руководитель", "ФИО, должность, степень, звание: ФИОНаучРук, ДолжНаучРук, СтепеньНаучРук, ЗваниеНаучРук."),
        ("ФИОДП", "ФИО студента в требуемом падеже для ВКР (обычно родительный)."),
        ("СегодняшняяДата", "Текущая дата формирования документов."),
    ]
    for name, desc in fields:
        doc.add_paragraph(f"• {name}: {desc}")

    doc.add_heading("Имена выходных файлов", level=1)
    doc.add_paragraph(
        "Имена документов формируются автоматически и включают ФИО и группу, "
        "например: «Дневник_{ФИО}_{Группа}.docx», «Титул_ВКР_{ФИО}_{Группа}.docx»."
    )

    doc.add_heading("Где используются данные", level=1)
    doc.add_paragraph(
        "Полный перечень соответствий «поле → документ(ы)» указан на странице сервиса ниже в инструкции. "
        "Заполняйте все поля без пропусков — пустые ячейки приводят к незаполненным местам в результатах."
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

@app.get("/instruction")
def download_instruction():
    """
    Отдаём инструкцию (DOCX). Если в корне лежит готовый файл (instruction.docx / инструкция.docx / …),
    вернём его. Иначе — сгенерируем типовой DOCX на лету.
    """
    for p in INSTRUCTION_CANDIDATES:
        if p.exists():
            return FileResponse(
                str(p),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                filename=INSTRUCTION_DOWNLOAD_NAME,
                headers={"Cache-Control": "no-store, no-cache, must-revalidate"},
            )
    # fallback: сгенерируем docx
    content = _build_instruction_docx_bytes()
    return StreamingResponse(
        io.BytesIO(content),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f'attachment; filename="{INSTRUCTION_DOWNLOAD_NAME}"',
            "Cache-Control": "no-store, no-cache, must-revalidate",
        },
    )

# ============= HTTP API =============
@app.get("/", response_class=HTMLResponse)
def index():
    return HTMLResponse(INDEX_HTML)

@app.post("/inspect")
def inspect(
    table_file: Optional[UploadFile] = File(default=None),
    gsheet_url: Optional[str] = Form(default=None),
    header_row: int = Form(default=1),
):
    # приоритет: если есть ссылка — используем её, иначе файл
    if gsheet_url and gsheet_url.strip():
        record, meta, cols = extract_record_from_gsheet(gsheet_url.strip(), header_row)
    elif table_file and (table_file.filename or "").strip():
        record, meta, cols = extract_record_from_upload(table_file, header_row)
    else:
        raise HTTPException(400, "Укажите Google Sheet ИЛИ выберите файл")

    needed = ["ФИО", "Группа"]
    missing = [k for k in needed if k not in record]

    if meta["mode"] == "wide":
        preview = record
        return JSONResponse({"columns": cols or [], "preview": preview, "missing": missing, "meta": meta})
    else:
        preview_pairs = list(record.items())[:12]
        return JSONResponse({"columns": [], "preview_pairs": preview_pairs, "missing": missing, "meta": meta})

@app.post("/generate")
def generate_zip(
    table_file: Optional[UploadFile] = File(default=None),
    gsheet_url: Optional[str] = Form(default=None),
    header_row: int = Form(default=1),
):
    if gsheet_url and gsheet_url.strip():
        record, meta, _ = extract_record_from_gsheet(gsheet_url.strip(), header_row)
    elif table_file and (table_file.filename or "").strip():
        record, meta, _ = extract_record_from_upload(table_file, header_row)
    else:
        raise HTTPException(400, "Укажите Google Sheet ИЛИ выберите файл")

    fio = safe(record.get("ФИО")) or "record"
    folder = slugify(f"001_{fio}")

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for tpl in TEMPLATES:
            try:
                # контекст: {tpl_key: значение из record по названию колонки}
                ctx = {tpl_key: safe(record.get(excel_col, "")) for tpl_key, excel_col in tpl["fields"].items()}
                doc = DocxTemplate(tpl["path"])
                doc.render(ctx)

                # имя файла из шаблонной маски out
                out_name = tpl["out"].format_map(SafeMap(record))
                out_name = slugify(out_name) or "doc_001.docx"

                out_mem = io.BytesIO(); doc.save(out_mem)
                zf.writestr(f"{folder}/{out_name}", out_mem.getvalue())
            except Exception as e:
                err = slugify(tpl.get("out","file")) + ".ERROR.txt"
                zf.writestr(f"{folder}/{err}", f"Ошибка ({tpl['path']}): {type(e).__name__}: {e}")

    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/zip",
        headers={"Content-Disposition": 'attachment; filename="generated_docs.zip"'}
    )

@app.get("/healthz")
def healthz():
    return PlainTextResponse("ok")
